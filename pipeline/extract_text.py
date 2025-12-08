from __future__ import annotations
import argparse
from pathlib import Path
from typing import Iterable

SUPPORTED_EXTS = {".pdf", ".docx"}


def iter_files(roots: Iterable[Path]) -> Iterable[Path]:
    for root in roots:
        for p in root.rglob("*"):
            if p.is_file() and p.suffix.lower() in SUPPORTED_EXTS:
                yield p


def extract_pdf(p: Path, max_pages: int | None = None) -> tuple[str, list[list[list[str]]]]:
    import pdfplumber
    text_parts: list[str] = []
    tables_all: list[list[list[str]]] = []
    with pdfplumber.open(str(p)) as pdf:
        n = len(pdf.pages)
        limit = min(n, max_pages) if max_pages is not None else n
        for i in range(limit):
            page = pdf.pages[i]
            t = page.extract_text(x_tolerance=1.5, y_tolerance=1.5) or ""
            text_parts.append(t.strip())
            # tables
            page_tables = page.extract_tables() or []
            for tbl in page_tables:
                normalized = [[(cell or "").strip() for cell in row] for row in tbl]
                if any(any(cell for cell in row) for row in normalized):
                    tables_all.append(normalized)
    return ("\n\n".join(filter(None, text_parts)), tables_all)


def extract_docx(p: Path) -> str:
    from docx import Document
    doc = Document(str(p))
    paras = [para.text.strip() for para in doc.paragraphs if para.text]
    return "\n".join(paras)


def write_text(out_path: Path, content: str, overwrite: bool) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    if out_path.exists() and not overwrite:
        return
    out_path.write_text(content, encoding="utf-8")


def write_tables_to_excel(out_path: Path, tables: list[list[list[str]]], overwrite: bool) -> None:
    if not tables:
        return
    out_path.parent.mkdir(parents=True, exist_ok=True)
    if out_path.exists() and not overwrite:
        return
    import pandas as pd
    with pd.ExcelWriter(out_path, engine="openpyxl") as writer:
        for idx, tbl in enumerate(tables, start=1):
            max_cols = max((len(r) for r in tbl), default=0)
            norm_rows = [r + [""] * (max_cols - len(r)) for r in tbl]
            df = pd.DataFrame(norm_rows)
            sheet_name = f"table_{idx}"[:31]
            df.to_excel(writer, index=False, header=False, sheet_name=sheet_name)


def main() -> int:
    ap = argparse.ArgumentParser(description="Bulk extract text (and tables) from PDFs/DOCX in folders 1 and 2")
    project_root = Path(__file__).parent.parent
    ap.add_argument("--roots", nargs="*", default=[
        str(project_root / "1"),
        str(project_root / "2"),
    ])
    ap.add_argument("--out", default=str(project_root / "outputs" / "raw"))
    ap.add_argument("--tables", action="store_true", help="Also extract tables to Excel files")
    ap.add_argument("--max-pages", type=int, default=None, help="Limit pages per PDF for faster runs")
    ap.add_argument("--overwrite", action="store_true")
    args = ap.parse_args()

    roots = [Path(p) for p in args.roots]
    out_root = Path(args.out)

    processed = 0
    errors = 0

    for src in iter_files(roots):
        try:
            rel = None
            # Find which root this belongs to for mirrored structure
            for r in roots:
                try:
                    rel = src.relative_to(r)
                    folder_tag = r.name  # "1" or "2"
                    break
                except Exception:
                    continue
            if rel is None:
                rel = src.name
                folder_tag = "unknown"

            base = src.stem
            dest_dir = out_root / folder_tag / rel.parent
            text_out = dest_dir / f"{base}.txt"

            if src.suffix.lower() == ".pdf":
                text, tables = extract_pdf(src, max_pages=args.max_pages)
                write_text(text_out, text, overwrite=args.overwrite)
                if args.tables:
                    excel_out = dest_dir / f"{base}_tables.xlsx"
                    write_tables_to_excel(excel_out, tables, overwrite=args.overwrite)
            elif src.suffix.lower() == ".docx":
                text = extract_docx(src)
                write_text(text_out, text, overwrite=args.overwrite)

            processed += 1
            print(f"✔ {src} -> {text_out}")
        except Exception as e:
            errors += 1
            print(f"✘ {src}: {e}")

    print(f"Done. Processed={processed}, Errors={errors}")
    return 0 if errors == 0 else 1


if __name__ == "__main__":
    raise SystemExit(main())
